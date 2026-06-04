#!/usr/bin/env python3
"""Render .diagnostics/weekly-intel/DIGEST.md to PDF + DOCX.

Used to preview what the weekly intel report looks like delivered as a
file (email attachment / WhatsApp document). The GitHub Issue remains
the canonical interactive surface (checkboxes are clickable there);
this is a read-only export for archival or out-of-band review.

Run:
    python tools/render_digest.py
    python tools/render_digest.py --in=path/to/DIGEST.md --out-dir=...
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem, HRFlowable,
)
from reportlab.lib.enums import TA_LEFT

ROOT = Path(__file__).resolve().parent.parent
DEFAULT_IN = ROOT / '.diagnostics' / 'weekly-intel' / 'DIGEST.md'
DEFAULT_OUT_DIR = ROOT / '.diagnostics' / 'weekly-intel'

# Strip the machine-readable trailers — useful in interactive issue
# bodies, distracting in a printable report.
TRAILER_RE = re.compile(r'<!--\s*intel:[^>]+?-->\s*')
DETAILS_RE = re.compile(r'</?details[^>]*>|<summary[^>]*>|</summary>', re.I)


def clean_line(line: str) -> str:
    line = TRAILER_RE.sub('', line)
    line = DETAILS_RE.sub('', line).rstrip()
    return line


def parse_inline(text: str) -> str:
    """Markdown → minimal HTML for reportlab Paragraph rendering.
    Order matters: bold before italic, links last."""
    text = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', text)
    text = re.sub(r'(?<!\w)\*(.+?)\*(?!\w)', r'<i>\1</i>', text)
    text = re.sub(r'`([^`]+?)`', r'<font face="Courier" color="#5b3aa8">\1</font>', text)
    text = re.sub(
        r'\[([^\]]+?)\]\(([^)]+?)\)',
        r'<link href="\2" color="#1a0f5c"><u>\1</u></link>',
        text,
    )
    # Lone <url> → clickable
    text = re.sub(
        r'&lt;(https?://[^\s&]+?)&gt;',
        r'<link href="\1" color="#1a0f5c"><u>\1</u></link>',
        text,
    )
    return text


# ── PDF renderer (reportlab) ─────────────────────────────────────────────


def render_pdf(md_text: str, out_path: Path) -> None:
    doc = SimpleDocTemplate(
        str(out_path), pagesize=A4,
        leftMargin=18 * mm, rightMargin=18 * mm,
        topMargin=18 * mm, bottomMargin=18 * mm,
        title='NAC — Weekly Intel Digest',
        author='NAC Brochure intel pipeline',
    )

    base = getSampleStyleSheet()
    styles = {
        'h1': ParagraphStyle('h1', parent=base['Heading1'],
                             fontSize=20, leading=24,
                             textColor=colors.HexColor('#1a0f5c'),
                             spaceAfter=10),
        'h2': ParagraphStyle('h2', parent=base['Heading2'],
                             fontSize=15, leading=20,
                             textColor=colors.HexColor('#5b3aa8'),
                             spaceBefore=14, spaceAfter=6),
        'h3': ParagraphStyle('h3', parent=base['Heading3'],
                             fontSize=12, leading=16,
                             textColor=colors.HexColor('#1a0f5c'),
                             spaceBefore=10, spaceAfter=4),
        'body': ParagraphStyle('body', parent=base['BodyText'],
                               fontSize=10, leading=14, alignment=TA_LEFT),
        'bullet': ParagraphStyle('bullet', parent=base['BodyText'],
                                 fontSize=10, leading=13,
                                 leftIndent=14, bulletIndent=2),
        'check_open': ParagraphStyle('co', parent=base['BodyText'],
                                     fontSize=10, leading=14,
                                     leftIndent=14, bulletIndent=2),
    }

    flow = []
    lines = [clean_line(l) for l in md_text.splitlines()]

    bullet_buf: list[Paragraph] = []

    def flush_bullets():
        nonlocal bullet_buf
        if bullet_buf:
            flow.append(ListFlowable(
                [ListItem(p, leftIndent=14) for p in bullet_buf],
                bulletType='bullet', bulletColor=colors.HexColor('#5b3aa8'),
                start='•', leftIndent=14,
            ))
            bullet_buf = []

    for raw in lines:
        if not raw.strip():
            flush_bullets()
            flow.append(Spacer(1, 4))
            continue
        if raw.startswith('# '):
            flush_bullets()
            flow.append(Paragraph(parse_inline(raw[2:]), styles['h1']))
            continue
        if raw.startswith('## '):
            flush_bullets()
            flow.append(Paragraph(parse_inline(raw[3:]), styles['h2']))
            continue
        if raw.startswith('### '):
            flush_bullets()
            flow.append(Paragraph(parse_inline(raw[4:]), styles['h3']))
            continue
        if raw.startswith('---'):
            flush_bullets()
            flow.append(Spacer(1, 4))
            flow.append(HRFlowable(width='100%', thickness=0.5,
                                   color=colors.HexColor('#bbb')))
            flow.append(Spacer(1, 4))
            continue
        m = re.match(r'^(\s*)- \[( |x)\] (.*)', raw)
        if m:
            flush_bullets()
            indent, mark, body = m.groups()
            box = '☑' if mark == 'x' else '☐'
            flow.append(Paragraph(
                f'<font color="#5b3aa8" size="12">{box}</font> ' + parse_inline(body),
                styles['check_open'],
            ))
            continue
        m = re.match(r'^(\s*)- (.*)', raw)
        if m:
            bullet_buf.append(Paragraph(parse_inline(m.group(2)), styles['bullet']))
            continue
        # Plain paragraph
        flush_bullets()
        flow.append(Paragraph(parse_inline(raw), styles['body']))

    flush_bullets()
    doc.build(flow)


# ── DOCX renderer (python-docx) ──────────────────────────────────────────


def add_runs(p, text: str) -> None:
    """Render inline **bold**, *italic*, `code`, links as runs."""
    pat = re.compile(
        r'(\*\*(.+?)\*\*'
        r'|(?<!\w)\*(.+?)\*(?!\w)'
        r'|`([^`]+?)`'
        r'|\[([^\]]+?)\]\(([^)]+?)\)'
        r'|<(https?://[^\s>]+?)>'
        r')'
    )
    i = 0
    for m in pat.finditer(text):
        if m.start() > i:
            p.add_run(text[i:m.start()])
        bold, italic, code, link_text, link_url, bare_url = (
            m.group(2), m.group(3), m.group(4),
            m.group(5), m.group(6), m.group(7),
        )
        if bold is not None:
            r = p.add_run(bold)
            r.bold = True
        elif italic is not None:
            r = p.add_run(italic)
            r.italic = True
        elif code is not None:
            r = p.add_run(code)
            r.font.name = 'Courier New'
            r.font.color.rgb = RGBColor(0x5b, 0x3a, 0xa8)
        elif link_text is not None:
            r = p.add_run(link_text)
            r.font.color.rgb = RGBColor(0x1a, 0x0f, 0x5c)
            r.underline = True
        elif bare_url is not None:
            r = p.add_run(bare_url)
            r.font.color.rgb = RGBColor(0x1a, 0x0f, 0x5c)
            r.underline = True
        i = m.end()
    if i < len(text):
        p.add_run(text[i:])


def render_docx(md_text: str, out_path: Path) -> None:
    doc = Document()

    # Base font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)

    for raw in md_text.splitlines():
        line = clean_line(raw)
        if not line.strip():
            doc.add_paragraph('')
            continue
        if line.startswith('# '):
            h = doc.add_heading(level=0)
            r = h.add_run(line[2:])
            r.font.color.rgb = RGBColor(0x1a, 0x0f, 0x5c)
            continue
        if line.startswith('## '):
            h = doc.add_heading(level=1)
            r = h.add_run(line[3:])
            r.font.color.rgb = RGBColor(0x5b, 0x3a, 0xa8)
            continue
        if line.startswith('### '):
            h = doc.add_heading(level=2)
            r = h.add_run(line[4:])
            r.font.color.rgb = RGBColor(0x1a, 0x0f, 0x5c)
            continue
        if line.startswith('---'):
            p = doc.add_paragraph()
            p.add_run('─' * 60).font.color.rgb = RGBColor(0xbb, 0xbb, 0xbb)
            continue
        m = re.match(r'^(\s*)- \[( |x)\] (.*)', line)
        if m:
            _, mark, body = m.groups()
            p = doc.add_paragraph(style='List Bullet')
            box = '☑ ' if mark == 'x' else '☐ '
            r = p.add_run(box)
            r.font.color.rgb = RGBColor(0x5b, 0x3a, 0xa8)
            add_runs(p, body)
            continue
        m = re.match(r'^(\s*)- (.*)', line)
        if m:
            p = doc.add_paragraph(style='List Bullet')
            add_runs(p, m.group(2))
            continue
        p = doc.add_paragraph()
        add_runs(p, line)

    doc.save(str(out_path))


# ── Main ─────────────────────────────────────────────────────────────────


def main() -> int:
    in_path = DEFAULT_IN
    out_dir = DEFAULT_OUT_DIR
    for a in sys.argv[1:]:
        if a.startswith('--in='):
            in_path = Path(a.split('=', 1)[1])
        elif a.startswith('--out-dir='):
            out_dir = Path(a.split('=', 1)[1])

    if not in_path.exists():
        sys.exit(f'❌ digest not found: {in_path}')
    out_dir.mkdir(parents=True, exist_ok=True)

    md = in_path.read_text(encoding='utf-8')
    pdf_out = out_dir / 'DIGEST.pdf'
    docx_out = out_dir / 'DIGEST.docx'
    render_pdf(md, pdf_out)
    render_docx(md, docx_out)
    print(f'Wrote {pdf_out.relative_to(ROOT)} ({pdf_out.stat().st_size:,} bytes)')
    print(f'Wrote {docx_out.relative_to(ROOT)} ({docx_out.stat().st_size:,} bytes)')
    return 0


if __name__ == '__main__':
    sys.exit(main())
